from abc import ABC, abstractmethod

from docx import Document
from docx.shared import Inches, RGBColor

from co_razon.model.judge import Case, Fact

DESCRIPTION = "Descripción: "
LIST_BULLET = "List Bullet"


class ReportGenerator(ABC):
    def __init__(self, case: Case):
        self._case = case

    @abstractmethod
    def generate_report(self, file_path: str):
        pass


class DocxReportGenerator(ReportGenerator):

    def __init__(self, case: Case):
        super(DocxReportGenerator, self).__init__(case)

    def generate_report(self, file_path: str):
        doc: Document = Document()

        # HEADING
        doc.add_heading(f"INFORME CASO: {self._case.name}", 0)

        # CASE HYPOTHESIS
        self.__hypothesis_report(doc)

        # FAVORABLE FACTS
        p_fav_facts = doc.add_paragraph()
        fav_run = p_fav_facts.add_run("Hechos favorables")
        fav_run.bold = True
        fav_run.font.color.rgb = RGBColor(128, 152, 72)

        if len(self._case.pretense.fav_facts) > 0:
            for num, fact in enumerate(self._case.pretense.fav_facts, start=1):
                self.__fact_report(doc, fact, num, 1)
        else:
            p = doc.add_paragraph("NINGUNO")
            p.paragraph_format.left_indent = Inches(0.2)

        # UNFAVORABLE FACTS
        p_unfav_facts = doc.add_paragraph()
        unfav_run = p_unfav_facts.add_run("Hechos desfavorables")
        unfav_run.bold = True
        unfav_run.font.color.rgb = RGBColor(236, 5, 142)

        if len(self._case.pretense.unfav_facts) > 0:
            for num, fact in enumerate(self._case.pretense.unfav_facts, start=1):
                self.__fact_report(doc, fact, f"{num}", 1)
        else:
            p = doc.add_paragraph("NINGUNO")
            p.paragraph_format.left_indent = Inches(0.2)

        # SAVE DOCUMENT
        doc.save(file_path)

    def __fact_report(self, doc: Document, fact: Fact, num: str, level: int):
        p_fact = doc.add_paragraph()
        p_fact_run = p_fact.add_run(f"{num}. {fact.label}")
        p_fact_run.bold = True
        p_fact.paragraph_format.left_indent = Inches(0.2 * level)

        p_fact_desc = doc.add_paragraph(style=LIST_BULLET)
        p_fact_desc.paragraph_format.left_indent = Inches(0.5 * level)
        p_fact_desc.add_run(DESCRIPTION).italic = True
        p_fact_desc.add_run(fact.desc)

        p_fact_desc = doc.add_paragraph(style=LIST_BULLET)
        p_fact_desc.paragraph_format.left_indent = Inches(0.5 * level)
        p_fact_desc.add_run("Relevancia: ").italic = True
        p_fact_desc.add_run(fact.relevance)

        p_fact_desc = doc.add_paragraph(style=LIST_BULLET)
        p_fact_desc.paragraph_format.left_indent = Inches(0.5 * level)
        p_fact_desc.add_run("Peso probatorio: ").italic = True
        p_fact_desc.add_run(fact.probatory_weight)

        p_fact_desc = doc.add_paragraph(style=LIST_BULLET)
        p_fact_desc.paragraph_format.left_indent = Inches(0.5 * level)
        p_fact_desc.add_run("Regla de la experiencia: ").italic = True
        p_fact_desc.add_run(fact.generate_experience_rule())

        # SUB HECHOS FAVORABLES
        if len(fact.fav_facts) > 0:
            p_fav_fact = doc.add_paragraph()
            fav_fact_run = p_fav_fact.add_run("Subhechos favorables")
            fav_fact_run.bold = True
            fav_fact_run.font.color.rgb = RGBColor(98, 187, 193)
            p_fav_fact.paragraph_format.left_indent = Inches(0.2 * level)

            for f_num, sub_fact in enumerate(fact.fav_facts, start=1):
                self.__fact_report(doc, sub_fact, f"{num}.{f_num}", level + 1)

        # SUB HECHOS DESFAVORABLES
        if len(fact.unfav_facts) > 0:
            p_unfav_fact = doc.add_paragraph()
            unfav_fact_run = p_unfav_fact.add_run("Subhechos desfavorables")
            unfav_fact_run.bold = True
            unfav_fact_run.font.color.rgb = RGBColor(98, 187, 193)
            p_unfav_fact.paragraph_format.left_indent = Inches(0.2 * level)

            for f_num, sub_fact in enumerate(fact.unfav_facts, start=1):
                self.__fact_report(doc, sub_fact, f"{num}.{f_num}", level + 1)

        # PRUEBAS FAVORABLES
        p_fav_evid = doc.add_paragraph()
        fav_run = p_fav_evid.add_run("Pruebas favorables")
        fav_run.bold = True
        fav_run.font.color.rgb = RGBColor(98, 187, 193)
        p_fav_evid.paragraph_format.left_indent = Inches(0.2 * level)

        if len(fact.fav_evidence) > 0:
            for e_num, evidence in enumerate(fact.fav_evidence, start=1):
                self.__evidence_report(doc, evidence, num, e_num, level)
        else:
            p = doc.add_paragraph("NINGUNA")
            p.paragraph_format.left_indent = Inches(0.4 * level)

        # PRUEBAS DESFAVORABLES
        p_fav_evid = doc.add_paragraph()
        fav_run = p_fav_evid.add_run("Pruebas desfavorables")
        fav_run.bold = True
        fav_run.font.color.rgb = RGBColor(245, 138, 7)
        p_fav_evid.paragraph_format.left_indent = Inches(0.2 * level)

        if len(fact.unfav_evidence) > 0:
            for e_num, evidence in enumerate(fact.unfav_evidence, start=1):
                self.__evidence_report(doc, evidence, num, e_num, level)
        else:
            p = doc.add_paragraph("NINGUNA")
            p.paragraph_format.left_indent = Inches(0.4 * level)


    def __evidence_report(self, doc: Document, evidence, fact_num: int, num: int, level: int):
        p_evidence = doc.add_paragraph()
        p_ev_run = p_evidence.add_run(f"{fact_num}.{num} {evidence.label}")
        p_ev_run.bold = True
        p_evidence.paragraph_format.left_indent = Inches(0.4 * level)

        p_ev_desc = doc.add_paragraph(style=LIST_BULLET)
        p_ev_desc.paragraph_format.left_indent = Inches(0.7 * level)
        p_ev_desc.add_run("Tipo de prueba: ").italic = True
        p_ev_desc.add_run(evidence.type)

        p_ev_desc = doc.add_paragraph(style=LIST_BULLET)
        p_ev_desc.paragraph_format.left_indent = Inches(0.7 * level)
        p_ev_desc.add_run(DESCRIPTION).italic = True
        p_ev_desc.add_run(evidence.desc)

        p_ev_desc = doc.add_paragraph(style=LIST_BULLET)
        p_ev_desc.paragraph_format.left_indent = Inches(0.7 * level)
        p_ev_desc.add_run("Credibilidad: ").italic = True
        p_ev_desc.add_run(evidence.credibility)

        p_ev_desc = doc.add_paragraph(style=LIST_BULLET)
        p_ev_desc.paragraph_format.left_indent = Inches(0.7 * level)
        p_ev_desc.add_run("Pertinencia: ").italic = True
        p_ev_desc.add_run(evidence.relevance)

        p_ev_desc = doc.add_paragraph(style=LIST_BULLET)
        p_ev_desc.paragraph_format.left_indent = Inches(0.7 * level)
        p_ev_desc.add_run("Regla de la experiencia: ").italic = True
        p_ev_desc.add_run(evidence.generate_experience_rule())

    def __hypothesis_report(self, doc: Document):
        p_hypothesis = doc.add_paragraph()
        p_hypothesis.add_run(f"Pretensión: {self._case.pretense.label}").bold = True
        p_hy_desc = doc.add_paragraph(style=LIST_BULLET)
        p_hy_desc.add_run(DESCRIPTION).italic = True
        p_hy_desc.add_run(self._case.pretense.desc)

        p_hy_desc = doc.add_paragraph(style=LIST_BULLET)
        p_hy_desc.add_run("Peso probatorio: ").italic = True
        p_hy_desc.add_run(self._case.pretense.probatory_weight)
